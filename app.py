# from flask import Flask, render_template, request, jsonify, session, send_file
# import requests
# from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
# from reportlab.lib.styles import getSampleStyleSheet
# from docx import Document
# from io import BytesIO

# app = Flask(__name__)
# app.secret_key = "cinema_secret"

# OLLAMA_API = "http://localhost:11434/api/generate"
# MODEL = "granite4:micro"


# # ---------- AI FUNCTION ----------
# # def generate_ai(prompt):
# #     payload = {
# #         "model": MODEL,
# #         "prompt": prompt,
# #         "stream": False,
# #         "temperature": 0.7
# #     }

# #     res = requests.post(OLLAMA_API, json=payload)
# #     return res.json().get("response", "")



# # def generate_ai(prompt, temperature=0.7):
# #     payload = {
# #         "model": MODEL,
# #         "prompt": prompt,
# #         "stream": False,
# #         "temperature": temperature
# #     }

# #     try:
# #         res = requests.post(OLLAMA_API, json=payload, timeout=120)
# #         return res.json().get("response", "").strip()
# #     except Exception as e:
# #         print("AI Error:", e)
# #         return "AI generation failed. Please try again."
    
# def generate_ai(prompt, temperature=0.7):
#     payload = {
#         "model": MODEL,
#         "prompt": prompt,
#         "stream": False,
#         "temperature": temperature
#     }

#     try:
#         res = requests.post(OLLAMA_API, json=payload, timeout=180)

#         if res.status_code != 200:
#             print("Ollama error:", res.text)
#             return "AI generation failed."

#         data = res.json()
#         return data.get("response", "").strip()

#     except Exception as e:
#         print("AI ERROR:", e)
#         return "AI generation failed."

# generated_data = {
#     "screenplay": "",
#     "characters": "",
#     "sound": ""
# }


# # ---------- ROUTES ----------

# @app.route("/")
# def landing():
#     return render_template("landing.html")


# @app.route("/set_username", methods=["POST"])
# def set_username():
#     session['username'] = request.json['username']
#     return jsonify(success=True)


# @app.route("/dashboard")
# def dashboard():
#     return render_template("dashboard.html", username=session.get("username", "Guest"))


# # @app.route("/generate", methods=["POST"])
# # def generate():
# #     story = request.json['story']

# #     screenplay_prompt = f"Write professional screenplay:\n{story}"
# #     char_prompt = f"Create detailed character profiles:\n{story}"
# #     sound_prompt = f"Suggest cinematic sound design:\n{story}"

# #     generated_data["screenplay"] = generate_ai(screenplay_prompt)
# #     generated_data["characters"] = generate_ai(char_prompt)
# #     generated_data["sound"] = generate_ai(sound_prompt)

# #     return jsonify(success=True)

# @app.route("/generate", methods=["POST"])
# def generate():
#     story = request.json['story']
#     genre = request.json.get("genre", "cinematic drama")

#     # ---------- SCREENPLAY PROMPT ----------
#     screenplay_prompt = f"""
# You are an award-winning Hollywood screenwriter and visual storyteller.

# Generate a professional cinematic screenplay based on this story.

# GENRE: {genre}

# STRICT FORMAT:
# • Scene headings in ALL CAPS
# • Format: INT./EXT. LOCATION – TIME
# • Character names centered above dialogue
# • Concise cinematic action lines
# • Strong dramatic pacing
# • Include 4–6 scenes

# FOR EACH SCENE ALSO PROVIDE:
# • Suggested image description (for AI image generation)
# • Suggested video clip idea (camera angle/movement)
# • Visual mood/colors
# • Suggested UI animation style (fade, parallax, slide-in, zoom)

# OUTPUT STRUCTURE:

# TITLE

# SCENE 1:
# Heading
# Action
# Dialogue
# Visual Suggestion:
# Video Suggestion:
# Animation Hint:

# SCENE 2:
# ...

# STORY:
# {story}
# """

#     # ---------- CHARACTER PROMPT ----------
#     character_prompt = f"""
# You are a professional cinematic character designer and visual consultant.

# Create 3 cinematic characters based on this story.

# For each character include:

# • Name
# • Age
# • Short background
# • Personality traits
# • Main motivation
# • Character arc (brief)
# • Physical appearance description (for image generation)
# • Costume/style notes
# • Suggested intro animation (slide-in, glow, dramatic reveal)

# Keep response concise but cinematic.

# Story:
# {story}
# """


#     # ---------- SOUND DESIGN PROMPT ----------
#     sound_prompt = f"""
# You are a Hollywood film sound designer and multimedia director.

# Create a scene-wise sound and media design plan.

# For each scene include:

# • Background music style
# • Ambient sounds
# • Foley effects
# • Emotional tone
# • Suggested related video imagery
# • Suggested image mood references
# • UI animation cues synced with sound

# Keep it cinematic and production-ready.

# Story:
# {story}
# """ 

#     # ---------- GENERATION ----------
#     generated_data["screenplay"] = generate_ai(screenplay_prompt, 0.75)
#     generated_data["characters"] = generate_ai(character_prompt, 0.65)
#     generated_data["sound"] = generate_ai(sound_prompt, 0.6)

#     return jsonify(success=True)



# @app.route("/screenplay")
# def screenplay():
#     return render_template("screenplay.html",
#                            content=generated_data["screenplay"])




# @app.route("/characters")
# def characters():
#     return render_template("characters.html", content=generated_data["characters"])


# @app.route("/sound")
# def sound():
#     return render_template("sound.html", content=generated_data["sound"])


# # ---------- DOWNLOAD TXT ----------
# @app.route("/download/txt")
# def txt():
#     buf = BytesIO()
#     buf.write(generated_data["screenplay"].encode())
#     buf.seek(0)
#     return send_file(buf, as_attachment=True, download_name="script.txt")


# # ---------- DOWNLOAD PDF ----------
# @app.route("/download/pdf")
# def pdf():
#     buf = BytesIO()
#     doc = SimpleDocTemplate(buf)
#     styles = getSampleStyleSheet()

#     screenplay_text = generated_data["screenplay"]

#     elements = [
#         Paragraph("Screenplay", styles['Title']),
#         Spacer(1, 12),
#         Paragraph(screenplay_text.replace("\n", "<br/>"),
#                   styles['Normal'])
#     ]

#     doc.build(elements)
#     buf.seek(0)

#     return send_file(buf,
#                      as_attachment=True,
#                      download_name="script.pdf")



# # ---------- DOWNLOAD DOCX ----------
# @app.route("/download/docx")
# def docx():
#     buf = BytesIO()
#     doc = Document()

#     doc.add_heading("Screenplay", level=1)
#     doc.add_paragraph(generated_data["screenplay"])

#     doc.save(buf)
#     buf.seek(0)

#     return send_file(buf,
#                      as_attachment=True,
#                      download_name="script.docx")

# if __name__ == "__main__":
#     print("Starting Flask server...")
#     app.run(host="127.0.0.1", port=5000, debug=True)


from flask import Flask, render_template, request, jsonify, session, send_file
import requests
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from io import BytesIO
import os
import re
from flask import jsonify

app = Flask(__name__)
app.secret_key = "cinema_secret"

OLLAMA_API = "http://localhost:11434/api/generate" #local Ollama server API endpoint
MODEL = "granite4:micro"

#25 free credits - one time character generation will generate 3 images - left credits will be 5.5 
#image generation function
def generate_character_image(prompt):

    headers = {
        "authorization": f"Bearer {STABILITY_API_KEY}",
        "accept": "image/*"
    }

    files = {
        "prompt": (None, prompt),
        "output_format": (None, "png")
    }

    response = requests.post(STABILITY_URL, headers=headers, files=files)

    if response.status_code != 200:
        print("Stability API Error:", response.text)
        return None

    os.makedirs("static/generated", exist_ok=True)

    filename = f"static/generated/char_{len(os.listdir('static/generated'))}.png"

    with open(filename, "wb") as f:
        f.write(response.content)

    return filename

# ---------- AI FUNCTION ----------
def generate_ai(prompt, temperature=0.7):
    payload = {
        "model": MODEL,
        "prompt": prompt,
        "stream": False,
        "temperature": temperature
    }

    try:
        res = requests.post(OLLAMA_API, json=payload, timeout=180)

        if res.status_code != 200:
            print("Ollama error:", res.text)
            return "AI generation failed."

        data = res.json()
        return data.get("response", "").strip()

    except Exception as e:
        print("AI ERROR:", e)
        return "AI generation failed."


# Global storage (single-user local app)
generated_data = {
    "screenplay": "",
    "characters": "",
    "sound": ""
}


def extract_scenes(screenplay):
    return re.findall(
        r"(SCENE\s*\d+:.*?)(?=SCENE\s*\d+:|$)",
        screenplay,
        re.DOTALL | re.IGNORECASE
    )

# ---------- ROUTES ----------

@app.route("/")
def landing():
    return render_template("landing.html")


@app.route("/set_username", methods=["POST"])
def set_username():
    session['username'] = request.json['username']
    return jsonify(success=True)


@app.route("/dashboard")
def dashboard():
    return render_template(
        "dashboard.html",
        username=session.get("username", "Guest")
    )


# ---------- GENERATE AI CONTENT ----------
@app.route("/generate", methods=["POST"])
def generate():
    story = request.json['story']
    genre = request.json.get("genre", "cinematic drama")

    # ---------- SCREENPLAY PROMPT ----------
    screenplay_prompt = f"""
You are an award-winning Hollywood screenwriter and visual storyteller.

Generate a professional cinematic screenplay based on this story.

GENRE: {genre}

STRICT FORMAT:
• Scene headings in ALL CAPS
• Format: INT./EXT. LOCATION – TIME
• Character names centered above dialogue
• Concise cinematic action lines
• Strong dramatic pacing
• Include 4–6 scenes must be included

FOR EACH SCENE ALSO PROVIDE:
• Suggested image description
• Suggested video clip idea
• Visual mood/colors
• Suggested UI animation style

OUTPUT STRUCTURE:

TITLE

SCENE :
Heading
Action
Dialogue
Visual Suggestion:
Video Suggestion:
Animation Hint:

STORY:
{story}
"""

    # ---------- CHARACTER PROMPT ----------
    character_prompt = f"""
You are a cinematic character designer.

Create 3 cinematic characters.

Include:
• Name
• Background
• Personality
• Motivation
• Character arc
• Appearance for image generation
• Costume/style
• Intro animation suggestion

Story:
{story}
"""

    # ---------- SOUND PROMPT ----------
    sound_prompt = f"""
You are a Hollywood sound designer.

Create cinematic sound design plan.

Include:
• Background music style
• Ambient sounds
• Foley effects
• Emotional tone
• Video imagery suggestions
• Image mood references
• UI animation sync ideas

Story:
{story}
"""

    # ---------- AI GENERATION ----------
    screenplay = generate_ai(screenplay_prompt, 0.75)
    characters = generate_ai(character_prompt, 0.65)
    sound = generate_ai(sound_prompt, 0.6)

    # Clean formatting (remove markdown artifacts)
    generated_data["screenplay"] = screenplay.replace("**", "").strip()
    generated_data["characters"] = characters.replace("**", "").strip()
    generated_data["sound"] = sound.replace("**", "").strip()

    return jsonify(success=True)


@app.route("/screenplay")
def screenplay():
    return render_template(
        "screenplay.html",
        content=generated_data["screenplay"]
    )


@app.route("/characters")
def characters():
    return render_template(
        "characters.html",
        content=generated_data["characters"]
    )


@app.route("/sound")
def sound():
    return render_template(
        "sound.html",
        content=generated_data["sound"]
    )


# ---------- DOWNLOAD TXT ----------
@app.route("/download/txt")
def txt():
    buf = BytesIO()
    buf.write(generated_data["screenplay"].encode())
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="script.txt")


# ---------- DOWNLOAD PDF ----------
@app.route("/download/pdf")
def pdf():
    buf = BytesIO()
    doc = SimpleDocTemplate(buf)
    styles = getSampleStyleSheet()

    elements = [
        Paragraph("Screenplay", styles['Title']),
        Spacer(1, 12)
    ]

    for line in generated_data["screenplay"].split("\n"):
        elements.append(Paragraph(line, styles['Normal']))

    doc.build(elements)
    buf.seek(0)

    return send_file(buf, as_attachment=True,
                     download_name="script.pdf")


# ---------- DOWNLOAD DOCX ----------
@app.route("/download/docx")
def docx():
    buf = BytesIO()
    doc = Document()

    doc.add_heading("Screenplay", level=1)
    doc.add_paragraph(generated_data["screenplay"])

    doc.save(buf)
    buf.seek(0)

    return send_file(buf, as_attachment=True,
                     download_name="script.docx")
    
    

@app.route("/generate_character_images", methods=["POST"])
def generate_character_images():

    import re

    text = generated_data.get("characters", "")

    if not text:
        return {"images": []}

    # STRICT match for exact header
    appearances = re.findall(
        r"Appearance\s+for\s+Image\s+Generation:\s*(.*?)(?:\n\n|$)",
        text,
        re.DOTALL | re.IGNORECASE
    )

    print("Exact Appearances found:", len(appearances))

    image_paths = []

    for appearance in appearances:

        prompt_text = appearance.strip()

        prompt = f"""
        cinematic character portrait,
        {prompt_text},
        professional film still,
        dramatic lighting,
        ultra realistic,
        85mm lens,
        high detail
        """

        print("Generating image for:", prompt_text[:80])

        img = generate_character_image(prompt)

        if img:
            image_paths.append(img)

    print("Generated images:", image_paths)

    return {"images": image_paths}


if __name__ == "__main__":
    print("Starting Flask server...")
    app.run(host="127.0.0.1", port=5000, debug=True)
